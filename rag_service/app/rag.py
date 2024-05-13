import os
import dotenv
from docx import Document
from langchain import hub
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter

dotenv.load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

llm = ChatOpenAI(model="gpt-3.5-turbo-0125")

class SimpleDocument:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {}

class DocxLoader:
    def __init__(self, folder_path):
        self.folder_path = folder_path

    def load(self):
        documents = []
        for filename in os.listdir(self.folder_path):
            if filename.endswith(".docx"):
                path = os.path.join(self.folder_path, filename)
                doc = Document(path)
                text = '\n\n'.join([para.text for para in doc.paragraphs if para.text])
                documents.append(SimpleDocument(text))
        return documents

loader = DocxLoader(folder_path="documents")
docs = loader.load()

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
splits = text_splitter.split_documents(docs)

vectorstore = Chroma.from_documents(documents=splits, embedding=OpenAIEmbeddings())

retriever = vectorstore.as_retriever()
prompt = hub.pull("rlm/rag-prompt")

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

rag_chain = (
    {"context": retriever | format_docs, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)
